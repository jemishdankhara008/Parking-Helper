"""One-time script to create README.docx from README.md content."""
# Documentation helper that exports a polished Word version of the project README for submission or sharing.
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(cell)
    return t

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p

def main():
    doc = Document()
    # The content is generated procedurally so the DOCX can be rebuilt after README wording changes.
    doc.add_heading('Parking Helper', 0)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Real-Time Smart Parking Detection System').bold = True
    p.add_run(' | Cambrian College AIE1014 Capstone')
    doc.add_paragraph()

    doc.add_paragraph(
        'Parking Helper uses YOLO26 to detect vehicles in parking lot images and video feeds. '
        'It provides a web dashboard for real-time monitoring, single-image analysis, and historical analytics.'
    )
    doc.add_paragraph()

    # Features
    doc.add_heading('Features', level=1)
    features = [
        'Dashboard — Real-time parking lot status (empty/occupied spots)',
        'Analyse Image — Upload a photo and get vehicle detections with bounding boxes',
        'Analytics — Occupancy over time and hourly averages from historical data',
        'Live Monitor — Process video feeds (local file or YouTube) with ROI-based spot mapping',
        'CLAHE Enhancement — Improved visibility in low-light and snow conditions',
        'Polygon IoU — Accurate spot occupancy using Shapely geometry',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')
    doc.add_paragraph()

    # Tech Stack
    doc.add_heading('Tech Stack', level=1)
    add_table(doc, ['Component', 'Technology'], [
        ['Object Detection', 'YOLO26x (Ultralytics)'],
        ['Backend API', 'FastAPI, Uvicorn'],
        ['Web UI', 'Streamlit'],
        ['Computer Vision', 'OpenCV'],
        ['Geometry', 'Shapely'],
        ['Data', 'Pandas, Altair'],
        ['Runtime', 'Python 3.10+'],
    ])
    doc.add_paragraph()

    # Project Structure
    doc.add_heading('Project Structure', level=1)
    doc.add_paragraph(
        'parking-helper/\n'
        '├── api/\n'
        '│   └── app.py           # FastAPI backend: /health, /predict\n'
        '├── ui/\n'
        '│   └── app_ui.py        # Streamlit UI: Dashboard, Analyse, Analytics, About\n'
        '├── main/\n'
        '│   ├── main.py         # Real-time video monitor (ROI + Shapely)\n'
        '│   └── roi/            # ROI selector tools\n'
        '├── data/\n'
        '│   ├── parking_lots.csv     # Config: ParkingLotID, URL, ROI\n'
        '│   ├── parking_status.json  # Live status (Dashboard reads this)\n'
        '│   └── reporting/           # *_history.csv for Analytics\n'
        '├── run_api.bat         # Quick-start script for API\n'
        '├── requirements.txt\n'
        '└── README.md'
    )
    doc.add_paragraph()

    # Prerequisites
    doc.add_heading('Prerequisites', level=1)
    for item in ['Python 3.10+', 'pip (Python package manager)', 'CUDA (optional, for GPU acceleration)']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

    # Installation
    doc.add_heading('Installation', level=1)
    doc.add_paragraph('1. Clone or download the project and navigate to the project folder:')
    add_code_block(doc, 'cd parking-helper')
    doc.add_paragraph('2. Create and activate a virtual environment (recommended):')
    add_code_block(doc, 'python -m venv venv\nvenv\\Scripts\\activate    # Windows\n# source venv/bin/activate   # Linux/Mac')
    doc.add_paragraph('3. Install dependencies:')
    add_code_block(doc, 'pip install -r requirements.txt')
    doc.add_paragraph('4. First run: The YOLO model (yolo26x.pt) will be downloaded automatically if not found locally.')
    doc.add_paragraph()

    # How to Run
    doc.add_heading('How to Run', level=1)
    doc.add_heading('Option 1: Analyse Single Image (API + UI)', level=2)
    doc.add_paragraph('1. Start the API (Terminal 1):')
    add_code_block(doc, 'cd parking-helper\npython -m uvicorn api.app:app --host 0.0.0.0 --port 8000')
    doc.add_paragraph('Or double-click run_api.bat.')
    doc.add_paragraph('2. Start the UI (Terminal 2):')
    add_code_block(doc, 'cd parking-helper\nstreamlit run ui/app_ui.py')
    doc.add_paragraph('3. Open the Streamlit URL (typically http://localhost:8501).')
    doc.add_paragraph('4. Go to Analyse Image, upload a parking lot photo, and click Analyse Parking Lot.')

    doc.add_heading('Option 2: Real-Time Video Monitor', level=2)
    doc.add_paragraph('1. Create data/parking_lots.csv:')
    add_table(doc, ['ParkingLotID', 'URL', 'ROI'], [
        ['Lot11', 'path/to/video.mp4', 'lot11_roi.csv'],
        ['Lot2', 'https://youtube.com/...', 'lot2_roi.csv'],
    ])
    doc.add_paragraph('2. Add ROI CSV files (Point1_X, Point1_Y, Point2_X, etc.) to data/.')
    doc.add_paragraph('3. Run the monitor:')
    add_code_block(doc, 'cd parking-helper\npython -m main.main')
    doc.add_paragraph('4. The monitor writes parking_status.json and *_history.csv. Start the UI to view the Dashboard and Analytics.')
    doc.add_paragraph()

    # API Endpoints
    doc.add_heading('API Endpoints', level=1)
    add_table(doc, ['Method', 'Endpoint', 'Description'], [
        ['GET', '/', 'Welcome message'],
        ['GET', '/health', 'Health check, model status'],
        ['GET', '/info', 'Model metadata'],
        ['POST', '/predict', 'Upload image, get vehicle detections'],
    ])
    doc.add_paragraph('Swagger UI: http://localhost:8000/docs')
    doc.add_paragraph('ReDoc: http://localhost:8000/redoc')
    doc.add_paragraph()

    # Data Files
    doc.add_heading('Data Files', level=1)
    add_table(doc, ['File', 'Purpose'], [
        ['parking_status.json', 'Live lot status (Dashboard)'],
        ['parking_status.csv', 'Combined status log (all lots)'],
        ['reporting/*_history.csv', 'Time-series per lot (Analytics charts)'],
        ['parking_lots.csv', 'Monitor config: lot ID, video URL, ROI file'],
    ])
    doc.add_paragraph()

    # Model Details
    doc.add_heading('Model Details', level=1)
    for item in [
        'Model: YOLO26x (Ultralytics)',
        'Task: Vehicle Object Detection',
        'Classes: Car, Motorcycle, Bus, Truck (COCO IDs: 2, 3, 5, 7)',
        'Confidence Threshold: 0.10 (API) / 0.15 (Monitor)',
        'Resolution: 1920px',
    ]:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

    # Troubleshooting
    doc.add_heading('Troubleshooting', level=1)
    add_table(doc, ['Issue', 'Solution'], [
        ['ModuleNotFoundError: No module named api', 'Run commands from parking-helper/ folder'],
        ['API Offline in UI', 'Start API first: python -m uvicorn api.app:app --host 0.0.0.0 --port 8000'],
        ['Port 8000 in use', 'Kill process or use --port 8001'],
        ['Model not loading', 'Ensure internet for first download; check path'],
    ])
    doc.add_paragraph()

    # Credits
    doc.add_heading('Credits', level=1)
    doc.add_paragraph('Project: Parking Helper — AIE1014 Capstone')
    doc.add_paragraph('Institution: Cambrian College')
    doc.add_paragraph('Technologies: YOLO26 (Ultralytics), FastAPI, OpenCV, Shapely, Streamlit')
    doc.add_paragraph()

    # License
    doc.add_heading('License', level=1)
    doc.add_paragraph('Capstone project — Cambrian College.')

    doc.save('README.docx')
    print('Created README.docx')

if __name__ == '__main__':
    main()
